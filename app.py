import base64
import configparser
import mimetypes
from io import BytesIO

import requests
from docx import Document
from docx.shared import Inches
from flask import Flask, render_template, request, send_file

app = Flask(__name__)

config = configparser.ConfigParser()
config.read('config.ini', encoding='utf-8')
GEMINI_API_KEY = config['Gemini']['API_KEY'].strip()
GEMINI_ENDPOINT = (
    'https://generativelanguage.googleapis.com/v1/models/'
    f'gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}'
)

PROMPT_TEMPLATE = '''請使用繁體中文，根據以下活動資訊，撰寫以下七個段落，請加上標題：

一、 主題說明  
二、 心得反思  
三、學習歷程檔案內容簡介（100 字內）  
四、檢討或反思  
五、學習或執行過程（步驟、結果）  
六、對未來的影響  
七、學習成果佐證說明  

活動資訊如下：
{user_input}
'''

def call_gemini_api(first_image_bytes, user_input, image_filename=None, timeout=60):
    '''
    與 Gemini 溝通：
    - 若有圖片：會連同第一張圖片與文字一起送出
    - 若無圖片：只送文字
    '''
    parts = []

    if first_image_bytes:
        mime = 'image/jpeg'
        if image_filename:
            guess = mimetypes.guess_type(image_filename)[0]
            if guess:
                mime = guess
        image_base64 = base64.b64encode(first_image_bytes).decode('utf-8')
        parts.append({
            'inlineData': {
                'mimeType': mime,
                'data': image_base64
            }
        })

    parts.append({'text': PROMPT_TEMPLATE.format(user_input=user_input)})

    payload = {'contents': [{'parts': parts}]}
    headers = {'Content-Type': 'application/json'}

    try:
        resp = requests.post(GEMINI_ENDPOINT, json=payload, headers=headers, timeout=timeout)
    except requests.RequestException as e:
        return f'[錯誤] 連線至 Gemini 失敗：{e}'

    if resp.status_code != 200:
        return f'[錯誤] Gemini 回應狀態碼 {resp.status_code}：{resp.text}'

    result = resp.json()
    try:
        text = result['candidates'][0]['content']['parts'][0]['text']
        return text
    except (KeyError, IndexError):
        return '[錯誤] Gemini 回傳內容解析失敗：\n' + str(result)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        image_files = request.files.getlist('images')
        user_input = request.form.get('user_input', '').strip()

        first_image_bytes = None
        first_image_name = None
        if image_files and image_files[0] and getattr(image_files[0], 'filename', ''):
            first_image_bytes = image_files[0].read()
            first_image_name = image_files[0].filename

        result_text = call_gemini_api(first_image_bytes, user_input, image_filename=first_image_name)

        doc = Document()
        doc.add_heading('學習歷程紀錄', 0)

        if '**' in result_text:
            for section in result_text.split('**'):
                s = section.strip()
                if s:
                    doc.add_paragraph(s)
        else:
            for line in result_text.splitlines():
                s = line.strip()
                if s:
                    doc.add_paragraph(s)

        doc.add_page_break()
        doc.add_heading('活動照片', level=1)

        for img in image_files:
            try:
                img.seek(0)
            except Exception:
                pass
            if getattr(img, 'filename', ''):
                try:
                    doc.add_picture(img, width=Inches(4))
                    doc.add_paragraph('')
                except Exception as e:
                    doc.add_paragraph(f'[提醒] 無法插入圖片 {img.filename}：{e}')

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name='learning.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    return render_template('index.html', generated_text=None)


if __name__ == '__main__':
    app.run(debug=True)